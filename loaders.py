import os
from langchain_core.documents import Document
from pypdf import PdfReader
import docx
import pandas as pd

def load_pdf(file_path: str) -> list[Document]:
    docs = []
    try:
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": file_path, "page": i + 1, "type": "pdf"}
                ))
    except Exception as e:
        print(f"Error leyendo PDF {file_path}: {e}")
    return docs

def load_docx(file_path: str) -> list[Document]:
    docs = []
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        text = "\n".join(full_text)
        if text:
            docs.append(Document(
                page_content=text,
                metadata={"source": file_path, "type": "docx"}
            ))
    except Exception as e:
        print(f"Error leyendo DOCX {file_path}: {e}")
    return docs

def load_excel(file_path: str) -> list[Document]:
    docs = []
    try:
        xls = pd.ExcelFile(file_path)
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            text_content = df.to_string(index=False)
            if text_content.strip():
                docs.append(Document(
                    page_content=f"Hoja: {sheet_name}\n\n{text_content}",
                    metadata={"source": file_path, "sheet": sheet_name, "type": "xlsx"}
                ))
    except Exception as e:
        print(f"Error leyendo Excel {file_path}: {e}")
    return docs

def process_file(file_path: str) -> list[Document]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return load_pdf(file_path)
    elif ext == '.docx':
        return load_docx(file_path)
    elif ext in ['.xlsx', '.xls']:
        return load_excel(file_path)
    else:
        print(f"Formato no soportado: {ext}")
        return []
